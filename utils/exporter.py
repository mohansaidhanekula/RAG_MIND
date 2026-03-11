"""
Export Module: Save Q&A sessions to PDF or Word (.docx).
"""

import io
from datetime import datetime
from typing import List, Dict


# ── PDF Export ─────────────────────────────────────────────────────────────────

def export_to_pdf(
    chat_history: List[Dict],
    doc_sources: List[str],
    title: str = "RAG Q&A Session",
) -> bytes:
    """
    Convert chat history to a formatted PDF.
    Returns raw PDF bytes.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm,  bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "Title2",
        parent=styles["Title"],
        fontSize=20,
        spaceAfter=6,
        textColor=colors.HexColor("#1e293b"),
    ))
    styles.add(ParagraphStyle(
        "Question",
        parent=styles["Normal"],
        fontSize=11,
        leading=16,
        textColor=colors.HexColor("#1d4ed8"),
        spaceBefore=14,
        spaceAfter=4,
        fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "Answer",
        parent=styles["Normal"],
        fontSize=10,
        leading=15,
        textColor=colors.HexColor("#1e293b"),
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        "Meta",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.grey,
    ))

    story = []

    # Header
    story.append(Paragraph(title, styles["Title2"]))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}",
        styles["Meta"]
    ))
    if doc_sources:
        story.append(Paragraph(
            f"Sources: {', '.join(doc_sources)}",
            styles["Meta"]
        ))
    story.append(Spacer(1, 0.4*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0")))
    story.append(Spacer(1, 0.4*cm))

    # Q&A pairs
    for i, turn in enumerate(chat_history, 1):
        story.append(Paragraph(f"Q{i}: {turn['question']}", styles["Question"]))
        answer_safe = turn["answer"].replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(answer_safe, styles["Answer"]))

        # Source citations
        if turn.get("sources"):
            src_names = list({
                s.metadata.get("source", "Unknown") for s in turn["sources"]
            })
            story.append(Paragraph(
                f"📎 Sources: {', '.join(src_names)}",
                styles["Meta"]
            ))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#f1f5f9")))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ── Word Export ────────────────────────────────────────────────────────────────

def export_to_docx(
    chat_history: List[Dict],
    doc_sources: List[str],
    title: str = "RAG Q&A Session",
) -> bytes:
    """
    Convert chat history to a .docx Word document.
    Returns raw docx bytes.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    t = doc.add_heading(title, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t.runs[0].font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

    # Meta
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}").font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    if doc_sources:
        sp = doc.add_paragraph()
        sp.add_run(f"Sources: {', '.join(doc_sources)}").font.size = Pt(9)
        sp.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    doc.add_paragraph()

    # Q&A
    for i, turn in enumerate(chat_history, 1):
        # Question
        q = doc.add_paragraph()
        qr = q.add_run(f"Q{i}: {turn['question']}")
        qr.bold = True
        qr.font.size = Pt(11)
        qr.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)

        # Answer
        a = doc.add_paragraph()
        ar = a.add_run(turn["answer"])
        ar.font.size = Pt(10)

        # Sources
        if turn.get("sources"):
            src_names = list({
                s.metadata.get("source", "Unknown") for s in turn["sources"]
            })
            sp = doc.add_paragraph()
            spr = sp.add_run(f"📎 Sources: {', '.join(src_names)}")
            spr.font.size = Pt(9)
            spr.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        doc.add_paragraph("─" * 60)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
